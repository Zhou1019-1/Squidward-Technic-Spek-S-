# -*- coding: utf-8 -*-
"""生成《章鱼频谱查看器 使用说明书.docx》。"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(BASE, "docs", "章鱼频谱查看器-使用说明书.docx")

PURPLE = RGBColor(0x5B, 0x3F, 0xA8)
GRAY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# 全局中文字体
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def set_cn(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color


def heading(text, level):
    h = doc.add_heading("", level=level)
    r = h.add_run(text)
    set_cn(r, bold=True, color=PURPLE if level <= 1 else None)


def para(text, size=11, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn(r, size=size, bold=bold, color=color)
    if align:
        p.alignment = align
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    set_cn(p.add_run(text))


# ============ 封面 ============
doc.add_paragraph()
icon = os.path.join(BASE, "icon", "HMSicon.png")
if os.path.exists(icon):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(icon, width=Cm(3.5))

para("章鱼频谱查看器", size=28, bold=True, color=PURPLE, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Octopus Spectrum Viewer", size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
para("—— 章鱼出品，必属精品 ——", size=12, color=PURPLE, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para("使 用 说 明 书", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para("版本 1.0.0 ｜ 2026 年 8 月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ============ 一、软件简介 ============
heading("一、软件简介", 1)
para(
    "章鱼频谱查看器是一款现代化的声学频谱分析软件，其算法设计致敬经典开源项目 "
    "Spek 及其活跃分支 Spek-X，并使用 Python + PySide6（Qt6）对界面与交互进行了全面重写与美化。"
)
para("主要用途：")
bullet("鉴别“假无损”音频：有损编码（MP3/AAC 等）通常会在高频处齐齐截断，"
       "转码成 FLAC/WAV 后截断痕迹依然保留，在频谱图上一目了然；")
bullet("观察音频的频率分布、动态范围与底噪特征；")
bullet("查看超声区（20kHz 以上）的隐藏信息与图案；")
bullet("导出精美频谱图用于分享与存档。")
para(
    "软件基于 FFmpeg 解码，几乎支持所有常见音频格式："
    "WAV、FLAC、MP3、M4A(AAC)、OGG、OPUS、APE、WMA、AIFF 等。"
)

# ============ 二、安装与启动 ============
heading("二、安装与启动", 1)
heading("2.1 安装程序（推荐）", 2)
bullet("下载 章鱼频谱查看器-setup-1.0.0.exe；")
bullet("双击运行，按向导完成安装（可选择创建桌面快捷方式）；")
bullet("从开始菜单或桌面快捷方式启动；安装包自带卸载程序。")
heading("2.2 绿色单文件", 2)
bullet("下载 章鱼频谱查看器.exe，免安装，双击直接运行。")
heading("2.3 源码运行", 2)
para("需要 Python 3.10+，执行以下命令：")
p = doc.add_paragraph()
r = p.add_run("pip install PySide6 numpy av\npython run.py")
set_cn(r, size=10)
r.font.name = "Consolas"
para("也支持命令行直接打开文件：python run.py 音频文件路径")

# ============ 三、界面介绍 ============
heading("三、界面介绍", 1)
shot = os.path.join(BASE, "docs", "screenshot-main.png")
if os.path.exists(shot):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(shot, width=Cm(16))
    para("图 1  主界面", size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
para("界面自上而下分为五个区域：")
bullet("标题栏：软件图标与名称，右侧为最小化 / 最大化 / 关闭按钮；")
bullet("工具栏：打开文件、调色板、窗函数、FFT 大小、声道选择、导出 PNG；")
bullet("信息面板：当前文件的编码格式、码率、采样率、位深、声道数、时长；")
bullet("频谱区：横轴为时间，纵轴为频率，颜色深浅代表该时刻该频率的能量（dB）；"
       "左侧为频率刻度尺，下方为时间刻度尺，右侧为调色板色条与 dB 刻度；")
bullet("状态栏：左侧为鼠标悬停读数，右侧为当前频率显示范围。")

# ============ 四、基本操作 ============
heading("四、基本操作", 1)
heading("4.1 打开音频文件", 2)
bullet("方式一：点击工具栏「🐙 打开文件」按钮选择文件；")
bullet("方式二：将音频文件直接拖入窗口。")
para("打开后软件立即开始分析，频谱图像会逐列“生长”，无需等待全曲分析完毕。")
heading("4.2 读取频谱信息", 2)
bullet("将鼠标移到频谱区，出现十字线与读数框，实时显示该点的 时间 / 频率 / 电平(dB)；")
bullet("底部状态栏同步显示相同读数。")
heading("4.3 缩放频率轴", 2)
bullet("在频谱区滚动鼠标滚轮：以光标位置为中心放大 / 缩小频率显示范围；")
bullet("双击频谱区：复位为全频段。")
para("该功能用于聚焦观察超声区（20kHz 以上）的隐藏内容。", color=GRAY)
heading("4.4 调整分析参数", 2)
bullet("调色板：章鱼 Octopus（默认）/ SoX 经典 / 光谱 Spectrum / 灰度 Mono；")
bullet("窗函数：Hann（默认）/ Hamming / Blackman-Harris；")
bullet("FFT 大小：2⁸ ~ 2¹⁴（默认 2¹¹ = 2048），数值越大频率分辨率越高、时间分辨率越低；")
bullet("声道：混合（默认）或单独查看某一声道。")
para("除调色板即时生效外，其余参数切换后会自动重新计算。", color=GRAY)
heading("4.5 导出频谱图", 2)
bullet("点击「导出 PNG」，选择保存位置即可，导出内容与屏幕所见一致。")
shot2 = os.path.join(BASE, "docs", "screenshot-sox-zoom.png")
if os.path.exists(shot2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(shot2, width=Cm(16))
    para("图 2  SoX 经典调色板 + 频率轴缩放（11~22kHz）", size=9, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

# ============ 五、如何鉴别假无损 ============
heading("五、如何鉴别“假无损”", 1)
para("有损编码为了节省码率会砍掉人耳不敏感的高频成分，典型截断频率：")
bullet("MP3 128kbps：约 16kHz 截断；")
bullet("MP3 320kbps：约 20kHz 截断；")
bullet("AAC 128kbps：约 15~16kHz 截断。")
para(
    "真正的无损音乐（CD 抓轨）频谱通常能自然延伸到 20kHz 以上。"
    "若一个 FLAC/WAV 文件的频谱在 16kHz 处出现整齐的“一刀切”，"
    "则很可能是由 MP3 等有损格式转码而来。"
)
para("提示：使用滚轮缩放到 15~22kHz 区间，截断痕迹会更加明显。", color=GRAY)

# ============ 六、技术说明 ============
heading("六、技术说明", 1)
bullet("解码层：FFmpeg (PyAV)，统一归一化为 float32；")
bullet("分析流水线：列数 = 绘图区像素宽度（像素即数据），时间区间按误差累积均匀分摊，"
       "每个像素列内做若干次加窗 FFT 并对幅度取平均；")
bullet("FFT：默认 2048 点、1025 频带，hop = 窗长（不重叠），dB 范围 0 ~ -120；")
bullet("渲染：Qt6 QPainter 自绘，NumPy 向量化着色，自适应刻度尺（1/2/5 因子表）；")
bullet("界面框架：PySide6 (Qt6)，无边框深色主题。")

# ============ 七、常见问题 ============
heading("七、常见问题", 1)
para("Q1：为什么打开某些文件提示“无法打开文件”？", bold=True)
para("A：请确认文件未损坏且为音频文件；极少数特殊编码可能不被支持。")
para("Q2：频谱图顶部为什么是空的？", bold=True)
para("A：说明该音频在对应频率以上没有能量，常见于有损压缩文件（高频被编码器切除）。")
para("Q3：FFT 大小调大后为什么声音细节反而“糊”了？", bold=True)
para("A：FFT 窗越长，频率越精细但时间越模糊（时频不确定性原理），请按需要在两者之间权衡。")
para("Q4：软件支持哪些操作系统？", bold=True)
para("A：安装包适用于 Windows 10/11 (64 位)；源码可在 Windows / macOS / Linux 运行。")

doc.add_paragraph()
para("—— 章鱼出品，必属精品 ——", size=10, color=PURPLE, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
with open(os.path.join(BASE, "tests", "docx_log.txt"), "w", encoding="utf-8") as f:
    f.write(f"saved: {OUT} ({os.path.getsize(OUT) // 1024} KB)")
